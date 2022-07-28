import wikipedia
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import Mm
import requests
import io
from docx.shared import Inches

name = input("Introdu numele tau: ")
wikipedia.set_lang("ro")
hs = input("La ce liceu esti?\n")
cls = input("In ce clasa esti?\n")
date = input("Pe ce data trebuie facut proiectul?\n")
title = input("Despre ce vrei sa fie proiectul tau?\n")
while True:
    try:
        wiki = wikipedia.page(title)
        break
    except:
        print("Nume proiect invalid")
        title = input("Introdu alt nume de proiect: \n")
text = wiki.content
text = re.sub(r'==', '', text)
text = re.sub(r'=', '', text)
text = re.sub(r'\n', '\n    ', text)
split = text.split('Vezi și', 1)
text = split[0]
print(text)

document = Document()

section = document.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)




paragraph = document.add_paragraph(date)
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
paragraph = document.add_paragraph(name)
paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph = document.add_paragraph('Clasa '+cls)
paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph = document.add_paragraph(hs)
paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
paragraph = document.add_heading(title, 0)
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph('    ' + text)
paragraph.style = document.styles['Normal']
paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_wiki_image(document, image_url):
    response = requests.get(image_url, stream=True)
    image = io.BytesIO(response.content)
    try:
        document.add_picture(image, width=Inches(1.5))
    except:
        pass

for image_url in wiki.images:
    add_wiki_image(document, image_url)


document.save(title+"- "+name + ".docx")
input()